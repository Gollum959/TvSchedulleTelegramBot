import pendulum
import re
import os
import shutil
import docx
import telebot


today = pendulum.today()

list_dir = os.listdir(path="E:\Work\Python\Project\Test Directory")
this_week_local_file_name = "E:\Work\Python\Project\Test Right Directory\\thisweek.docx"
next_week_local_file_name = "E:\Work\Python\Project\Test Right Directory\\nextweek.docx"

current_week_start = today.start_of('week')
current_week_end = today.end_of('week')
next_week_start = current_week_start.add(weeks=1)
next_week_end = current_week_end.add(weeks=1)


def regular_string(week_start, week_end):
    regular_string = r'(0' + str(week_start.day) + '|' + str(week_start.day) + ')\
[.]*\
(0' + str(week_start.month) + '|' + str(week_start.month) + ')*\
[.]*\
(' + str(week_start.year) + ')*\
[\s]*-[\s]*\
(0' + str(week_end.day) + '|' + str(week_end.day) + ')\
[.]*\
(0' + str(week_end.month) + '|' + str(week_end.month) + ')*\
[.]*\
(' + str(week_end.year) + ')*'
    return regular_string


def path(week_start, week_end):
    regular_strig_week = regular_string(week_start, week_end)
    for directory in list_dir:
        match_dir = re.findall(regular_strig_week, directory)
        if match_dir:
            return f'E:\Work\Python\Project\Test Directory\{directory}'
    return ''


def file_path(week_start, week_end):
    path_file = path(week_start, week_end)
    if path_file:
        list_file = os.listdir(path=path_file)
        for file in list_file:
            match_file = re.findall(r"(p|р)(т|t)(s)\w*(.doc|.docx|rtf)$", file, re.IGNORECASE)
            if match_file:
                return f'{path_file}\{file}'
    return ''


def chek_and_copy(file_path, local_file, text):
    sourse_file_info = os.stat(file_path)
    sourse_file_size = sourse_file_info.st_size

    if os.path.isfile(local_file):
        local_file_info = os.stat(local_file)
        local_file_size = local_file_info.st_size
    else:
        local_file_size = 0

    if sourse_file_size == local_file_size:
        print(f'File {text} is already exist')
    else:
        shutil.copyfile(file_path, local_file, follow_symlinks=True)
        print(f'File {text} is copied')


def chek_file_exist(soure_file_path, local_file_path, text):
    if os.path.isfile(soure_file_path):
        chek_and_copy(soure_file_path, local_file_path, text)
    else:
        print(f'Sourse file {text} isn`t found')


this_weak_sourse_file_path = file_path(current_week_start, current_week_end)
next_weak_sourse_file_path = file_path(next_week_start, next_week_end)

if this_weak_sourse_file_path:
    chek_file_exist(this_weak_sourse_file_path, this_week_local_file_name, 'for this week')
else:
    print(f'Sourse file for this week isn`t found')

if next_weak_sourse_file_path:
    chek_file_exist(next_weak_sourse_file_path, next_week_local_file_name, 'for next week')
else:
    print(f'Sourse file for next week isn`t found')

def print_paragraps(text_in_cells):
    result = ''
    for date_cell_text in text_in_cells:
        if re.findall(r"\w", date_cell_text.text):
            result += remove_strike_text(date_cell_text)+' '
    return result

def remove_strike_text(paragraph):
    lst = paragraph.runs
    small = filter(lambda x: not x.font.strike, lst)
    result = ''
    for item in list(small):
        result += item.text
    if result:
        return result+'\n'
    else:
        return result


def show_day_timetable_pts(week_local_file_name, day_of_week, name_equipment):
    doc = docx.Document(week_local_file_name)
    tables = doc.tables
    result = print_paragraps(tables[List_PTS_and_hardware[name_equipment][0]].rows[day_of_week].cells[0].paragraphs)
    result += '_______________\n'
    work_schedule = print_paragraps(tables[List_PTS_and_hardware[name_equipment][0]].rows[day_of_week].cells[List_PTS_and_hardware[name_equipment][1]].paragraphs)
    if work_schedule.strip():
        result += work_schedule
    else:
        result += 'Работ нет'
    return result

def show_week_timetable_pts(week_local_file_name, name_equipment):
    doc = docx.Document(week_local_file_name)
    tables = doc.tables
    result = ''
    for i in range(1, 8):
        result += print_paragraps(tables[List_PTS_and_hardware[name_equipment][0]].rows[i].cells[0].paragraphs)
        result += ('__________\n')
        work_schedule = print_paragraps(tables[List_PTS_and_hardware[name_equipment][0]].rows[i].cells[List_PTS_and_hardware[name_equipment][1]].paragraphs)
        if work_schedule.strip().startswith('Изм.'):
            result += 'Работ нет\n'
        elif work_schedule.strip():
            result += work_schedule
        else:
            result += 'Работ нет\n'
        result += ('__________\n')
    return result


doc = docx.Document(this_week_local_file_name)
List_PTS_and_hardware = {}
for index_and_table in enumerate(doc.tables):
    number_of_colums = len(index_and_table[1].rows[0].cells)
    print (number_of_colums)
    for number in range(1, number_of_colums):
        name_equipment = index_and_table[1].rows[0].cells[number].text.strip()
        List_PTS_and_hardware[name_equipment] = [index_and_table[0], number]

print(List_PTS_and_hardware)


token = '2133524927:AAGL7Quqshq3OaLmfmFwTpSqaqWj25g2p0g'
bot = telebot.TeleBot(token)

@bot.message_handler(commands=['start'])
def handle_text(message):
    user_markup = telebot.types.ReplyKeyboardMarkup(True, False)
    for name_hardware in List_PTS_and_hardware:
        user_markup.row(name_hardware)
    bot.send_message(message.from_user.id, 'Добро пожаловать в бот по получению расписания ПТС!\n Выберите Пункт МЕНЮ', reply_markup=user_markup)

#@bot.callback_query_handlers(func=lambda call:True)
@bot.message_handler(func=lambda mess: mess.text, content_types=['text'])
def show_day_timetable(message):
    keyboard = telebot.types.InlineKeyboardMarkup()
    callback_data_menu_1 = {}
    callback_data_menu_1['id']=1
    callback_data_menu_1['message']=message.text
    menu_1 = telebot.types.InlineKeyboardButton(text='Расписание на сегодня', callback_data= "submenu_1_"+message.text)
    menu_2 = telebot.types.InlineKeyboardButton(text='Расписание на неделю', callback_data="submenu_2_"+message.text)
    menu_3 = telebot.types.InlineKeyboardButton(text='Расписание на Следующую неделю', callback_data="submenu_3_"+message.text)
    keyboard.add(menu_1, menu_2, menu_3)
    bot.send_message(message.chat.id, 'It works!', reply_markup=keyboard)

@bot.callback_query_handler(func=lambda call: call.data.startswith('submenu_'))
def show_day_timetable(call):
    splited_submenu = call.data.split('_')
    name_pts = splited_submenu[2]
    type_schedule = int(splited_submenu[1])
    if type_schedule == 1:
        result = show_day_timetable_pts(this_week_local_file_name, today.day_of_week, name_pts)
        bot.send_message(call.from_user.id, result)
    elif type_schedule == 2:
        result = show_week_timetable_pts(this_week_local_file_name, name_pts)
        bot.send_message(call.from_user.id, result)
    elif type_schedule == 3:
        result = show_week_timetable_pts(next_week_local_file_name, name_pts)
        bot.send_message(call.from_user.id, result)

bot.polling(none_stop=True)

